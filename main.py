import os
import shutil
import sys
import time
from io import BytesIO

import requests
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.text.run import Run
from tqdm import trange

img_count = 0


def get_img_urls():
    """
    读取文本中的图片编号
    :return: 返回图片编号生成的url列表
    """
    # 图片下载链接模板
    img_url_template = 'https://asia.pokemon-card.com/hk/card-img/hk%s.png'
    urls = []
    counts = []
    try:
        with open('./card.txt', 'r', encoding='utf-8') as file:
            # 读取文件内容
            lines = file.readlines()
            for line in lines:
                if line.startswith('#') or not line.strip():
                    continue
                line = line.strip()
                split = line.split(' ')
                id = split[0].zfill(8)
                urls.append(f"{img_url_template % id}")
                if len(split) != 2:
                    counts.append(4)
                else:
                    counts.append(int(split[1]))
    except:
        print('检查是否创建了card.txt文件，或者文件内容格式是否正确')
        raise Exception('检查是否创建了card.txt文件，或者文件内容格式是否正确')
    return urls, counts


def download_img_data(url: str):
    """
    下载图片数据
    :param url: 图片url地址
    :return: 图片的byte数据
    """
    try:
        res = requests.get(url)
        return BytesIO(res.content)
    except:
        print(f'图片下载失败:{url}')
        return None


def insert_img_into_doc(img_path: str, count: int, run: Run):
    """
    写入4张图片到word中
    :param img_path: 图片地址
    :param run: word文档的段落操作符
    :param count: 插入次数
    :return: None
    """
    global img_count
    for i in range(count):
        run.add_picture(img_path, Inches(3.5), Inches(2.5))
        img_count = img_count + 1
        if img_count % 2 == 0:
            run.add_break()
        else:
            run.add_text('   ')


def copy_docx(tmp_doc_name: str):
    """
    复制word文档模板
    :param tmp_doc_name: 新建的word名称
    :return: None
    """
    # 检查是否已打包成exe
    if getattr(sys, 'frozen', False):
        # 获取exe文件所在目录
        exe_dir = sys._MEIPASS if hasattr(sys, '_MEIPASS') else os.path.dirname(sys.executable)
        # 构建内部文件的完整路径
        source_path = os.path.join(exe_dir, 'template.docx')
    else:
        # 在开发环境中直接使用文件路径
        source_path = 'template.docx'
    shutil.copy(source_path, tmp_doc_name)


if __name__ == '__main__':
    try:
        # 获取文本中的图片id对应的下载地址列表
        img_urls, counts = get_img_urls()
        # 创建存储图片的临时文件夹
        os.makedirs('./tmp', exist_ok=True)
        # 从模板中复制一个新word，用时间命名
        tmp_doc_name = time.strftime('%Y%m%d%H%M%S', time.localtime()) + '.docx'
        copy_docx(tmp_doc_name)
        # 打开新word，获取当前段落的操作符
        doc = Document(tmp_doc_name)
        paragraph = doc.paragraphs[0]
        run = paragraph.add_run()
        for i in trange(len(img_urls), desc='图片下载中'):
            # 下载图片数据
            img_data = download_img_data(img_urls[i])
            if img_data:
                with Image.open(img_data) as img:
                    # 图片旋转90度
                    rotate_data = img.rotate(90, expand=True)
                    save_path = f'./tmp/{img_urls[i][-12:]}'
                    rotate_data.save(save_path)
                    # 写入到word中
                    insert_img_into_doc(save_path, counts[i], run)
        # 保存word文档
        doc.save(tmp_doc_name)
    finally:
        input('按任意键结束')
